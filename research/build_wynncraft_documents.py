"""Build polished DOCX editions of the Wynncraft research paper and guide.

The implementation follows the compact_reference_guide preset from the
Codex documents skill. The canonical source remains Markdown.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent

PAPER_SOURCE = ROOT / "wynncraft-endgame-class-building-paper.md"
GUIDE_SOURCE = ROOT / "wynncraft-endgame-class-building-guide.md"

PAPER_OUTPUT = ROOT / "Wynncraft-End-Game-Class-Building-Research-Paper.docx"
GUIDE_OUTPUT = ROOT / "Wynncraft-End-Game-Class-Building-Guide.docx"


@dataclass(frozen=True)
class DocumentSpec:
    source: Path
    output: Path
    kicker: str
    title: str
    subtitle: str
    description: str
    running_label: str


SPECS = (
    DocumentSpec(
        source=PAPER_SOURCE,
        output=PAPER_OUTPUT,
        kicker="Research Paper",
        title="End-Game Class Building in Wynncraft",
        subtitle="Mechanics, archetypes, Aspects, and optimization",
        description="A versioned research synthesis for Wynncraft 2.2.2",
        running_label="Wynncraft End-Game Class Building | Research Paper",
    ),
    DocumentSpec(
        source=GUIDE_SOURCE,
        output=GUIDE_OUTPUT,
        kicker="Operator Guide",
        title="Wynncraft End-Game Class Building Guide",
        subtitle="A practical workflow for version 2.2.2",
        description="Build contracts, class decisions, optimization targets, and validation",
        running_label="Wynncraft End-Game Class Building | Practical Guide",
    ),
)


# compact_reference_guide preset and named editorial-cover overrides.
INK = "172B3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66717A"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "AAB7C4"
CALLOUT = "F4F6F9"
CODE_FILL = "F2F4F7"
GOLD = "7A5A00"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def set_width_element(parent, tag: str, width: int, type_name: str = "dxa") -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), type_name)


def apply_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    set_width_element(tbl_pr, "w:tblW", CONTENT_WIDTH_DXA)
    set_width_element(tbl_pr, "w:tblInd", TABLE_INDENT_DXA)

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for existing in list(grid):
        grid.remove(existing)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            set_width_element(tc_pr, "w:tcW", width)
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

    set_table_borders(table)


def choose_column_widths(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> list[int]:
    column_count = len(headers)
    if column_count == 1:
        return [CONTENT_WIDTH_DXA]
    if column_count == 2:
        label_like = max((len(row[0]) for row in [headers, *rows]), default=0) <= 34
        return [2700, 6660] if label_like else [4680, 4680]
    if column_count == 3:
        return [2500, 2050, 4810]
    if column_count == 4:
        return [1760, 2500, 2700, 2400]
    if column_count == 5:
        return [1100, 1750, 1700, 2400, 2410]

    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section, label: str) -> None:
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label)
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run("Page ")
    set_run_font(label_run, size=9, color=MUTED)
    add_page_field(paragraph)


def set_style_font(style, name: str, size: float, color: str, bold: bool = False) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (11, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = document.styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    quote = document.styles["Quote"]
    set_style_font(quote, "Calibri", 11, DARK_BLUE, bold=False)
    quote.font.italic = True
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.25)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.25

    code = document.styles["No Spacing"]
    set_style_font(code, "Consolas", 9.5, INK)


def get_numbering_root(document: Document):
    return document.part.numbering_part.element


def next_numbering_id(numbering_root, tag: str, attr: str) -> int:
    values = []
    for element in numbering_root.findall(qn(tag)):
        value = element.get(qn(attr))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def create_numbering(document: Document, kind: str) -> int:
    root = get_numbering_root(document)
    abstract_id = next_numbering_id(root, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(root, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)

    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    root.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    if bold:
        bold_element = OxmlElement("w:b")
        r_pr.append(bold_element)
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|https?://[^\s)>]+|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)"
)


def add_inline(paragraph, text: str, *, base_size: float = 11) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=INK)
        token = match.group(0)
        if token.startswith("["):
            link_match = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
            else:
                run = paragraph.add_run(token)
                set_run_font(run, size=base_size, color=INK)
        elif token.startswith("http"):
            add_hyperlink(paragraph, token, token.rstrip(".,;"))
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(9, base_size - 1), color=DARK_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=INK, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=INK, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=INK)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def add_cover(document: Document, spec: DocumentSpec) -> None:
    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run(spec.kicker.upper())
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(spec.title)
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run(spec.subtitle)
    set_run_font(run, size=15, color=DARK_BLUE)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(46)
    run = description.add_run(spec.description)
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(4)
    run = metadata.add_run("Research baseline: Wynncraft 2.2.2 | 13 August 2026")
    set_run_font(run, size=10, color=INK, bold=True)

    authority = document.add_paragraph()
    authority.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authority.add_run("Live Wynncraft API authority with local solver implementation evidence")
    set_run_font(run, size=9.5, color=MUTED)

    document.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    if len(table_lines) < 2:
        return [], [], start

    def split_row(line: str) -> list[str]:
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    headers = split_row(table_lines[0])
    body = [split_row(line) for line in table_lines[2:]]
    return headers, body, index


def add_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    widths = choose_column_widths(headers, rows)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        set_run_font(run, size=9.2, color=INK, bold=True)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        for column_index in range(len(headers)):
            value = values[column_index] if column_index < len(values) else ""
            cell = row.cells[column_index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, base_size=8.8)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")

    apply_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def add_code_block(document: Document, code_lines: Iterable[str]) -> None:
    lines = list(code_lines)
    paragraph = document.add_paragraph(style="No Spacing")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, CODE_FILL)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9, color=INK)
        if index < len(lines) - 1:
            run.add_break()


def add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    cover_headings_to_skip = 2
    in_code = False
    code_lines: list[str] = []
    bullet_num_id: int | None = None
    decimal_num_id: int | None = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not stripped:
            bullet_num_id = None
            decimal_num_id = None
            index += 1
            continue

        if stripped.startswith("|"):
            headers, rows, next_index = parse_table(lines, index)
            if headers:
                add_table(document, headers, rows)
                index = next_index
                continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if cover_headings_to_skip > 0 and level <= 2:
                cover_headings_to_skip -= 1
                index += 1
                continue
            style_level = min(level - 1 if level > 1 else 1, 4)
            paragraph = document.add_paragraph(style=f"Heading {style_level}")
            run = paragraph.add_run(text.replace("`", ""))
            heading_values = {
                1: (16, BLUE),
                2: (16, BLUE),
                3: (13, BLUE),
                4: (12, DARK_BLUE),
            }
            size, color = heading_values[level]
            set_run_font(run, size=size, color=color, bold=True)
            bullet_num_id = None
            decimal_num_id = None
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph(style="Quote")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            item_text = bullet_match.group(1)
            if item_text.startswith("[ ] "):
                item_text = "☐ " + item_text[4:]
                paragraph.paragraph_format.left_indent = Inches(0.187)
            elif item_text.startswith("[x] ") or item_text.startswith("[X] "):
                item_text = "☒ " + item_text[4:]
                paragraph.paragraph_format.left_indent = Inches(0.187)
            else:
                if bullet_num_id is None:
                    bullet_num_id = create_numbering(document, "bullet")
                apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, item_text)
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            if decimal_num_id is None:
                decimal_num_id = create_numbering(document, "decimal")
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, decimal_num_id)
            add_inline(paragraph, number_match.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith("```")
                or next_line.startswith("> ")
                or re.match(r"^-\s+", next_line)
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            index += 1

        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))

    if in_code and code_lines:
        add_code_block(document, code_lines)


def set_core_properties(document: Document, spec: DocumentSpec) -> None:
    props = document.core_properties
    props.title = spec.title
    props.subject = spec.subtitle
    props.author = "OpenAI Codex"
    props.keywords = "Wynncraft, class building, archetypes, Aspects, optimization"
    props.comments = "Research baseline: Wynncraft 2.2.2, 13 August 2026"


def build_document(spec: DocumentSpec) -> None:
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    configure_page(section)
    configure_header_footer(section, spec.running_label)
    set_core_properties(document, spec)
    add_cover(document, spec)
    add_markdown(document, spec.source.read_text(encoding="utf-8"))

    for section in document.sections:
        configure_page(section)

    document.save(spec.output)


def main() -> None:
    for spec in SPECS:
        build_document(spec)
        print(spec.output)


if __name__ == "__main__":
    main()
